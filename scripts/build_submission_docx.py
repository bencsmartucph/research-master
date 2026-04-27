"""
build_submission_docx.py
Converts paper_draft_v3_final.md to a submission-quality AER-style Word document.
Usage: python scripts/build_submission_docx.py
"""
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

MANUSCRIPT = r"C:\Users\PKF715\Documents\claude_repos\Research_Master\manuscripts\paper_draft_v4_final.md"
OUTPUT     = r"C:\Users\PKF715\Documents\claude_repos\Research_Master\manuscripts\paper_final_submission.docx"
FONT       = "Times New Roman"

# ─── Inline markdown parser ───────────────────────────────────────────────────

def parse_inline(text):
    """Return list of (text, bold, italic, code) from markdown inline syntax."""
    tokens = []
    # Order matters: *** before ** before *
    pattern = re.compile(r'\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*|`([^`]+)`')
    last = 0
    for m in pattern.finditer(text):
        if m.start() > last:
            tokens.append((text[last:m.start()], False, False, False))
        if m.group(1):   tokens.append((m.group(1), True,  True,  False))
        elif m.group(2): tokens.append((m.group(2), True,  False, False))
        elif m.group(3): tokens.append((m.group(3), False, True,  False))
        elif m.group(4): tokens.append((m.group(4), False, False, True))
        last = m.end()
    if last < len(text):
        tokens.append((text[last:], False, False, False))
    return tokens

def add_inline(para, text, base_bold=False, base_italic=False, size=None):
    """Add formatted runs from markdown text to a paragraph."""
    for (t, bold, italic, code) in parse_inline(text):
        if not t:
            continue
        run = para.add_run(t)
        run.bold   = bold  or base_bold
        run.italic = italic or base_italic
        run.font.name = "Courier New" if code else FONT
        run.font.size = Pt(10 if code else (size or 12))

# ─── Paragraph formatting helpers ─────────────────────────────────────────────

def fmt_double(para):
    """Body paragraph spacing. Single-spacing with paragraph gap fits AER
    Insights page budget (~10-14 pages for 5500 words inc. refs and figures)."""
    pf = para.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_before = Pt(0)
    pf.space_after  = Pt(8)

def fmt_single(para, after=Pt(4)):
    pf = para.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_before = Pt(0)
    pf.space_after  = after

def add_page_number_field(para):
    """Insert a PAGE field into a paragraph."""
    run = para.add_run()
    for tag, attr in [('w:fldChar', {'w:fldCharType': 'begin'}),
                      ('w:instrText', None),
                      ('w:fldChar', {'w:fldCharType': 'end'})]:
        el = OxmlElement(tag)
        if attr:
            for k, v in attr.items():
                el.set(qn(k), v)
        if tag == 'w:instrText':
            el.set(qn('xml:space'), 'preserve')
            el.text = ' PAGE '
        run._r.append(el)

# ─── Table helpers ─────────────────────────────────────────────────────────────

def cell_borders_none(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'), 'none')
        borders.append(tag)
    tcPr.append(borders)

def cell_border_bottom(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ('top','left','right','insideH','insideV'):
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'), 'none')
        borders.append(tag)
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '8')
    bot.set(qn('w:color'), '000000')
    borders.append(bot)
    tcPr.append(borders)

def set_cell_shading(cell, fill):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill)
    tcPr.append(shd)

def add_md_table(doc, lines):
    """Parse markdown table lines and add to document."""
    rows = []
    for line in lines:
        stripped = line.strip()
        if re.match(r'^\|[\s\-:|]+\|$', stripped):
            continue  # separator row
        cells = [c.strip() for c in stripped.strip('|').split('|')]
        rows.append(cells)
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    rows  = [r + [''] * (ncols - len(r)) for r in rows]

    tbl = doc.add_table(rows=len(rows), cols=ncols)
    tbl.style = 'Table Grid'

    for i, row_data in enumerate(rows):
        row = tbl.rows[i]
        is_header = (i == 0)
        for j, raw in enumerate(row_data):
            cell = row.cells[j]
            p    = cell.paragraphs[0]
            p.paragraph_format.space_after  = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            # Strip bold markers for display, but detect them
            is_bold_cell = '**' in raw
            clean = re.sub(r'\*\*(.+?)\*\*', r'\1', raw)
            clean = re.sub(r'\*(.+?)\*', r'\1', clean)
            run = p.add_run(clean)
            run.font.name = FONT
            run.font.size = Pt(11)
            run.bold = is_bold_cell or is_header
            if is_header:
                set_cell_shading(cell, 'F2F2F2')
                cell_border_bottom(cell)
            else:
                cell_borders_none(cell)

# ─── Figure placeholder ────────────────────────────────────────────────────────

FIGURE_MAP = {
    '2': r'C:\Users\PKF715\Documents\claude_repos\Research_Master\outputs\figures\fig2_rti_vs_antiimmig_by_regime.png',
    '3': r'C:\Users\PKF715\Documents\claude_repos\Research_Master\outputs\figures\fig3_marginal_effects.png',
    '6': r'C:\Users\PKF715\Documents\claude_repos\Research_Master\outputs\figures\fig6_cwed_country_slopes.png',
}

def add_figure_placeholder(doc, text):
    """Insert figure image (if available) followed by italic caption.
    Falls back to grey placeholder if image not found."""
    import os
    inner = text.strip().lstrip('[').rstrip(']')
    inner = re.sub(r'\*\*(.+?)\*\*', r'\1', inner)
    inner = re.sub(r'\*(.+?)\*', r'\1', inner)

    # Extract figure number from "Figure N here:" pattern
    m = re.search(r'Figure\s+(\w+)\s+here', inner, re.IGNORECASE)
    fig_num = m.group(1) if m else None
    img_path = FIGURE_MAP.get(fig_num) if fig_num else None

    if img_path and os.path.exists(img_path):
        # Insert image
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fmt_single(p, after=Pt(2))
        p.paragraph_format.space_before = Pt(8)
        run = p.add_run()
        run.add_picture(img_path, width=Inches(5.5))

        # Insert caption (figure number bold, rest plain) under image
        caption_text = re.sub(r'^Figure\s+\w+\s+here:\s*', '', inner, flags=re.IGNORECASE)
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fmt_single(cap, after=Pt(8))
        bold_run = cap.add_run(f'Figure {fig_num}. ')
        bold_run.bold = True
        bold_run.font.name = FONT
        bold_run.font.size = Pt(10)
        cap_run = cap.add_run(caption_text)
        cap_run.italic = True
        cap_run.font.name = FONT
        cap_run.font.size = Pt(10)
        return

    # Fallback grey placeholder
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt_single(p, after=Pt(6))
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run(inner)
    run.italic = True
    run.font.name = FONT
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  'E8E8E8')
    pPr.append(shd)

# ─── Section heading helpers ────────────────────────────────────────────────────

def add_section_heading(doc, text):
    """Roman-numeral section heading: bold 12pt, space before/after."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(text)
    run.bold = True
    run.font.name = FONT
    run.font.size = Pt(12)

def add_subsection_heading(doc, text):
    """Letter subsection heading: bold italic 12pt."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(text)
    run.bold   = True
    run.italic = True
    run.font.name = FONT
    run.font.size = Pt(12)

# ─── Body paragraph ─────────────────────────────────────────────────────────────

def add_body_para(doc, text, first_para=False):
    """Double-spaced body paragraph. No first-line indent (AER style)."""
    p = doc.add_paragraph()
    fmt_double(p)
    # Small space before new paragraph to visually separate in double-spaced text
    p.paragraph_format.space_before = Pt(0)
    add_inline(p, text)
    return p

# ─── Abstract box ────────────────────────────────────────────────────────────────

def add_abstract(doc, text):
    """Single-spaced 11pt abstract with 0.5in indents."""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.left_indent  = Inches(0.5)
    p.paragraph_format.right_indent = Inches(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(12)
    add_inline(p, text, size=11)

# ─── Reference entry ─────────────────────────────────────────────────────────────

def add_reference(doc, text):
    """Single-spaced reference with 0.5in hanging indent, 10pt."""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.left_indent   = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)
    p.paragraph_format.space_before  = Pt(0)
    p.paragraph_format.space_after   = Pt(4)
    add_inline(p, text, size=10)

# ─── Main builder ────────────────────────────────────────────────────────────────

def build():
    with open(MANUSCRIPT, encoding='utf-8') as f:
        raw = f.read()

    lines = raw.split('\n')

    doc = Document()

    # Page setup
    for section in doc.sections:
        section.page_width   = Twips(12240)
        section.page_height  = Twips(15840)
        section.left_margin  = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin   = Inches(1)
        section.bottom_margin = Inches(1)

    # Default Normal style
    normal = doc.styles['Normal']
    normal.font.name = FONT
    normal.font.size = Pt(12)
    try:
        normal.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
    except Exception:
        pass

    # Footer: page number centred
    sec = doc.sections[0]
    footer_para = sec.footer.paragraphs[0]
    footer_para.clear()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number_field(footer_para)

    # ── Parse state machine ──────────────────────────────────────────────────
    in_abstract    = False
    in_references  = False
    abstract_lines = []
    table_buffer   = []
    in_table       = False

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # --- separators ---
        if stripped == '---':
            i += 1
            continue

        # --- blank lines ---
        if not stripped:
            if in_table and table_buffer:
                add_md_table(doc, table_buffer)
                table_buffer = []
                in_table = False
            i += 1
            continue

        # --- table rows ---
        if stripped.startswith('|') and '|' in stripped[1:]:
            in_table = True
            table_buffer.append(stripped)
            i += 1
            continue
        else:
            if in_table and table_buffer:
                add_md_table(doc, table_buffer)
                table_buffer = []
                in_table = False

        # --- title (# heading) ---
        if stripped.startswith('# ') and not stripped.startswith('## '):
            title_text = stripped[2:]
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(24)
            p.paragraph_format.space_after  = Pt(18)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            run = p.add_run(title_text)
            run.bold      = True
            run.font.name = FONT
            run.font.size = Pt(14)
            i += 1
            continue

        # --- author / affiliation / draft line (immediately after title) ---
        if stripped.startswith('**') and stripped.endswith('**') and len(stripped) < 60:
            name = stripped.strip('*')
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(name)
            run.bold      = False
            run.font.name = FONT
            run.font.size = Pt(12)
            i += 1
            continue

        # Lines that look like affiliation or draft info (not headings, not bold blocks)
        if not stripped.startswith('#') and not stripped.startswith('[') \
                and 'University of Copenhagen' in stripped:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(stripped)
            run.font.name = FONT
            run.font.size = Pt(12)
            i += 1
            continue

        if 'Welfare State Seminar' in stripped:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            p.paragraph_format.space_after = Pt(24)
            run = p.add_run("Welfare State Seminar — April 2026")
            run.italic    = True
            run.font.name = FONT
            run.font.size = Pt(11)
            i += 1
            continue

        # --- ## Abstract ---
        if stripped == '## Abstract':
            in_abstract = True
            # Abstract label
            p = doc.add_paragraph()
            p.paragraph_format.left_indent  = Inches(0.5)
            p.paragraph_format.right_indent = Inches(0.5)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after  = Pt(4)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            run = p.add_run("Abstract")
            run.bold      = True
            run.font.name = FONT
            run.font.size = Pt(11)
            i += 1
            continue

        if in_abstract and stripped:
            add_abstract(doc, stripped)
            in_abstract = False
            i += 1
            continue

        # --- ## References ---
        if stripped == '## References':
            in_references = True
            # Page break before references
            doc.add_paragraph().add_run().add_break(
                __import__('docx.enum.text', fromlist=['WD_BREAK']).WD_BREAK.PAGE)
            add_section_heading(doc, 'References')
            i += 1
            continue

        # Exit references mode when a new ## heading appears (e.g. Appendix)
        if in_references and stripped.startswith('## '):
            in_references = False

        if in_references and stripped:
            add_reference(doc, stripped)
            i += 1
            continue

        # --- ## Section headings (numbered) ---
        if stripped.startswith('## '):
            heading_text = stripped[3:]
            add_section_heading(doc, heading_text)
            i += 1
            continue

        # --- ### Subsection headings ---
        if stripped.startswith('### '):
            sub_text = stripped[4:]
            add_subsection_heading(doc, sub_text)
            i += 1
            continue

        # --- Figure placeholders ---
        if stripped.startswith('[') and 'Figure' in stripped and 'here' in stripped:
            add_figure_placeholder(doc, stripped)
            i += 1
            continue

        # --- Regular body paragraph ---
        if stripped:
            add_body_para(doc, stripped)

        i += 1

    # Flush any remaining table
    if table_buffer:
        add_md_table(doc, table_buffer)

    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")

if __name__ == '__main__':
    build()
